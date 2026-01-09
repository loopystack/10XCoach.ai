#!/usr/bin/env python3
"""Extract text from SBBE Book.docx"""
try:
    from docx import Document
    import sys
    
    doc_path = "SBBE Book FINAL 07 11 25 1752232058710-Alan_Wozniak_-_Epilogue.docx"
    
    try:
        doc = Document(doc_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Also try to get text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        full_text = "\n\n".join(text_parts)
        
        print("=" * 80)
        print(f"Extracted {len(full_text)} characters from {doc_path}")
        print("=" * 80)
        print(full_text)
        
    except FileNotFoundError:
        print(f"Error: File '{doc_path}' not found")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
        
except ImportError:
    print("Error: python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    print("Please run this script again.")
    sys.exit(1)

